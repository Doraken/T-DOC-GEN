import tkinter as tk
from tkinter import simpledialog
from docx import Document
import os
import json

def creer_fichier_word(nom_app, code_app, config):
    # Créer un document Word
    doc = Document()
    doc.add_heading('Informations de l\'Application', 0)
    doc.add_paragraph(f'Nom de l\'Application: {nom_app}')
    doc.add_paragraph(f'Code du Projet: {code_app}')
    
    # Ajouter les informations supplémentaires à partir du fichier de configuration
    for titre, contenu in config.items():
        doc.add_heading(titre, level=1)
        for sous_titre, details in contenu.items():
            doc.add_heading(sous_titre, level=2)
            for key, value in details.items():
                doc.add_paragraph(f'{key}: {value}')

    # Nommer le fichier en utilisant les entrées utilisateur
    output_dir = config.get('output', 'output')
    nom_fichier = f'DAT_{nom_app}_{code_app}.docx'
    output_path = os.path.join(output_dir, nom_fichier)
    doc.save(output_path)
    print(f'Fichier créé: {output_path}')

def demander_informations():
    root = tk.Tk()
    root.withdraw()  # Cacher la fenêtre Tkinter principale
    
    # Demander le nom de l'application et le code du projet
    nom_app = simpledialog.askstring("Nom de l'Application", "Entrez le nom de l'application :")
    code_app = simpledialog.askstring("Code du Projet", "Entrez le code du projet :")
    
    root.destroy()  # Fermer la fenêtre Tkinter après la saisie des informations
    return nom_app, code_app

def charger_configuration():
    # Charger les configurations à partir du fichier config.json
    config_file = os.path.join(os.path.dirname(__file__), '..', 'config', 'config_DAT.json')
    with open(config_file, 'r') as fichier:
        return json.load(fichier)

def run():
    nom_app, code_app = demander_informations()
    if nom_app and code_app:  # Vérifier que l'utilisateur a bien saisi les informations
        config = charger_configuration()
        creer_fichier_word(nom_app, code_app, config)
    else:
        print("Opération annulée.")

if __name__ == "__main__":
    run()
